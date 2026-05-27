from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "scinu_learning_ecosystem_report.md"
TARGET = BASE / "scinu_learning_ecosystem_report.docx"


def set_run_font(run, size=16, bold=False, color=None, font="TH Sarabun New"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font)


def set_para(paragraph, before=0, after=6, line=1.25, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text, style=None, size=16, bold=False, color=None, align=None, before=0, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    set_para(p, before=before, after=after, line=line, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    set_para(p, before=16 if level == 1 else 10, after=6, line=1.15)
    size = 20 if level == 1 else 17
    color = "2E74B5" if level == 1 else "1F4D78"
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def split_table_row(line):
    cells = [cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")]
    return cells


def add_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = [6.5 / col_count] * col_count
    for i, cell_text in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "F2F4F7")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=0, line=1.1)
        run = p.add_run(cell_text)
        set_run_font(run, size=14, bold=True, color="1F4D78")
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.15)
            run = p.add_run(cell_text)
            set_run_font(run, size=14)
    doc.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        if not re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", raw):
            rows.append(split_table_row(raw))
        i += 1
    return rows, i


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "TH Sarabun New"
    normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")

    title = lines[0].removeprefix("# ").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=96, after=18, line=1.15)
    r = p.add_run(title)
    set_run_font(r, size=24, bold=True, color="0B2545")

    i = 1
    while i < len(lines) and not lines[i].startswith("## "):
        text = lines[i].strip()
        if text:
            add_text(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
    doc.add_page_break()

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            i += 1 
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("- "):
            add_text(doc, line[2:].strip(), style="List Bullet", after=4)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            add_text(doc, re.sub(r"^\d+\.\s+", "", line), style="List Number", after=4)
            i += 1
            continue
        add_text(doc, line)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("NU Smart Admin Ecosystem")
    set_run_font(run, size=10, color="666666")
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build()
